import os
import re
import time
from datetime import datetime
from io import BytesIO
from typing import Optional

import anthropic
import httpx
from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import Response
from fpdf import FPDF
from pydantic import BaseModel
from sqlmodel import Session, select, text

from app.database import get_db
from app.dependencies import get_current_user
from app.models import (
    DocStatus,
    DocType,
    Document,
    DocumentVersion,
    Education,
    Experience,
    Job,
    JobEvent,
    JobEventType,
    Skill,
    UserSkill,
)

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
BUCKET = "documents"

# S3-004: Implement Document Upload Workflow
# Rules: S3-BR-004 (supported formats), S3-BR-005 (reject unsupported formats)
ALLOWED_RESUME_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "text/plain": ".txt",
}
MAX_RESUME_FILE_BYTES = 10 * 1024 * 1024  # 10 MB

router = APIRouter(prefix="/documents", tags=["documents"])


# ── Request models ────────────────────────────────────────────────────────────


class SaveCoverLetterRequest(BaseModel):
    job_id: Optional[str] = None
    cover_letter_text: str
    title: Optional[str] = None
    status: Optional[str] = "active"
    tags: Optional[str] = None


class GenerateCoverLetterRequest(BaseModel):
    job_id: str
    profile: dict


class ImproveCoverLetterRequest(BaseModel):
    job_id: str
    cover_letter_text: str
    instruction: str


class GenerateForJobRequest(BaseModel):
    job_id: str


class ImproveResumeRequest(BaseModel):
    resume_text: str
    instruction: str


class CreateVersionRequest(BaseModel):
    version_label: Optional[str] = None


class RestoreDocumentRequest(BaseModel):
    restore_to: Optional[str] = "draft"  # "draft" or "final"


class RenameDocumentRequest(BaseModel):
    title: str


# S3-009: Job-to-Library Linking
class LinkDocumentRequest(BaseModel):
    job_id: str
    document_type: DocType
    replace_existing: bool = False


class UnlinkDocumentRequest(BaseModel):
    job_id: str
    document_type: DocType


# ── Helpers ───────────────────────────────────────────────────────────────────


def generate_pdf(text: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(25, 25, 25)
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 7, sanitize_pdf_text(text))
    return bytes(pdf.output())


def sanitize_pdf_text(text: str) -> str:
    """FPDF's core fonts (helvetica, times, courier) only support Latin-1.
    AI-generated text often includes Unicode punctuation (em dashes, smart
    quotes) that crashes PDF generation. Replace common cases with ASCII
    equivalents rather than crashing the save."""
    replacements = {
        "\u2014": "-",  # em dash —
        "\u2013": "-",  # en dash –
        "\u2018": "'",  # left single quote '
        "\u2019": "'",  # right single quote '
        "\u201c": '"',  # left double quote "
        "\u201d": '"',  # right double quote "
        "\u2026": "...",  # ellipsis …
    }
    for unicode_char, ascii_equivalent in replacements.items():
        text = text.replace(unicode_char, ascii_equivalent)
    return text


EXPORT_CONTENT_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt": "text/plain",
}


def generate_docx_bytes(text: str) -> bytes:
    docx_file = DocxDocument()
    for paragraph in text.split("\n"):
        docx_file.add_paragraph(paragraph)
    buffer = BytesIO()
    docx_file.save(buffer)
    return buffer.getvalue()


def build_export_filename(document: Document, format: str) -> str:
    raw_name = document.title or document.file_name or "document"
    safe_name = re.sub(r"[^a-zA-Z0-9-_ ]", "", raw_name).strip().replace(" ", "_")
    return f"{safe_name or 'document'}.{format}"


def snapshot_document_version(db: Session, document: Document) -> None:
    """Save the current document state before replacing it."""
    current_version = document.version_number or 1
    db.add(
        DocumentVersion(
            document_id=document.id,
            user_id=document.user_id,
            version_number=current_version,
            version_label=document.version_label or f"v{current_version}",
            document_text=document.document_text,
            file_url=document.file_url,
        )
    )
    next_version = current_version + 1
    document.version_number = next_version
    document.version_label = f"v{next_version}"


def parse_document_status(status: Optional[str]) -> DocStatus:
    status_map = {
        "active": DocStatus.DRAFT,
        "draft": DocStatus.DRAFT,
        "final": DocStatus.FINAL,
        "archived": DocStatus.ARCHIVED,
    }
    document_status = status_map.get((status or "active").lower())
    if not document_status:
        raise HTTPException(status_code=400, detail="Invalid document status")
    return document_status


def clean_document_tags(tags: Optional[str]) -> Optional[str]:
    clean_tags = ",".join(tag.strip() for tag in (tags or "").split(",") if tag.strip())
    return clean_tags or None


def validate_resume_file(file: UploadFile, file_bytes: bytes) -> None:
    """Rules: S3-BR-004, S3-BR-005"""
    if file.content_type not in ALLOWED_RESUME_TYPES:
        raise HTTPException(
            status_code=400,
            detail=("Unsupported file type. Please upload a PDF, DOCX, or TXT file."),
        )
    if len(file_bytes) > MAX_RESUME_FILE_BYTES:
        raise HTTPException(
            status_code=400,
            detail="File is too large. Maximum size is 10 MB.",
        )
    if len(file_bytes) == 0:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")


@router.post("/{document_id}/versions")
def create_document_version(
    document_id: str,
    payload: CreateVersionRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    user_id = current_user.get("sub")

    doc = db.get(Document, document_id)
    if not doc or doc.user_id != user_id:
        raise HTTPException(status_code=404, detail="Document not found")

    # Calculate next version number
    next_version = (doc.version_number or 0) + 1

    # Snapshot current state into document_versions
    version = DocumentVersion(
        document_id=document_id,
        user_id=user_id,
        version_number=next_version,
        version_label=payload.version_label or f"v{next_version}",
        document_text=doc.document_text,
        file_url=doc.file_url,
    )
    db.add(version)

    # Update current version metadata on the document
    doc.version_number = next_version
    doc.version_label = version.version_label
    doc.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(version)

    return {
        "version_id": version.id,
        "version_number": version.version_number,
        "version_label": version.version_label,
        "created_at": version.created_at.isoformat(),
    }


@router.get("/{document_id}/versions")
def get_document_versions(
    document_id: str,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    user_id = current_user.get("sub")

    doc = db.get(Document, document_id)
    if not doc or doc.user_id != user_id:
        raise HTTPException(status_code=404, detail="Document not found")

    versions = db.exec(
        select(DocumentVersion)
        .where(DocumentVersion.document_id == document_id)
        .order_by(DocumentVersion.version_number.desc())
    ).all()

    return [
        {
            "version_id": v.id,
            "version_number": v.version_number,
            "version_label": v.version_label,
            "document_text": v.document_text,
            "file_url": v.file_url,
            "owner_id": v.user_id,
            "created_at": v.created_at.isoformat(),
        }
        for v in versions
    ]


async def upload_to_storage(
    file_bytes: bytes, path: str, content_type: str
) -> Optional[str]:  # noqa: E501
    async with httpx.AsyncClient() as client:
        upload_res = await client.post(
            f"{SUPABASE_URL}/storage/v1/object/{BUCKET}/{path}",
            headers={
                "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
                "Content-Type": content_type,
                "x-upsert": "true",
            },
            content=file_bytes,
        )
        if upload_res.status_code not in (200, 201):
            raise HTTPException(
                status_code=500, detail=f"Upload failed:{upload_res.text}"
            )  # noqa: E501

        sign_res = await client.post(
            f"{SUPABASE_URL}/storage/v1/object/sign/{BUCKET}/{path}",
            headers={
                "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
                "Content-Type": "application/json",
            },
            json={"expiresIn": 31536000},
        )
        sign_data = sign_res.json()
        signed_path = sign_data.get("signedURL") or sign_data.get("signedUrl", "")
        return f"{SUPABASE_URL}/storage/v1{signed_path}" if signed_path else None


# ── Resume endpoints ──────────────────────────────────────────────────────────


@router.get("/resume/job/{job_id}")
def get_resume_for_job(
    job_id: str,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    user_id = current_user.get("sub")
    record = db.exec(
        select(Document)
        .where(Document.job_id == job_id)
        .where(Document.user_id == user_id)
        .where(Document.doc_type == DocType.RESUME)
        .order_by(Document.updated_at.desc())
    ).first()

    if not record:
        raise HTTPException(status_code=404, detail="No resume found")

    return {
        "document_id": record.id,
        "resume_text": record.document_text,
        "file_name": record.file_name,
        "file_url": record.file_url,
    }


@router.post("/resume/save")
async def save_resume(
    file: UploadFile = File(...),
    resume_text: str = Form(...),
    job_id: Optional[str] = Form(None),
    title: Optional[str] = Form(None),
    status: Optional[str] = Form("active"),
    tags: Optional[str] = Form(None),
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    user_id = current_user.get("sub")
    linked_job_id = job_id or None

    if linked_job_id:
        job = db.exec(
            select(Job).where(Job.id == linked_job_id, Job.owner_id == user_id)
        ).first()
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")

    status_map = {
        "active": DocStatus.DRAFT,
        "draft": DocStatus.DRAFT,
        "final": DocStatus.FINAL,
        "archived": DocStatus.ARCHIVED,
    }
    document_status = status_map.get((status or "active").lower())
    if not document_status:
        raise HTTPException(status_code=400, detail="Invalid document status")

    clean_title = title.strip() if title else ""
    clean_tags = (
        ",".join(tag.strip() for tag in (tags or "").split(",") if tag.strip()) or None
    )

    file_bytes = await file.read()
    validate_resume_file(file, file_bytes)
    path = f"{user_id}/resume/{int(time.time())}_{file.filename}"

    file_url = await upload_to_storage(file_bytes, path, file.content_type)

    existing = None
    if linked_job_id:
        existing = db.exec(
            select(Document)
            .where(Document.job_id == linked_job_id)
            .where(Document.user_id == user_id)
            .where(Document.doc_type == DocType.RESUME)
        ).first()

    if existing:
        snapshot_document_version(db, existing)
        existing.title = clean_title or file.filename
        existing.status = document_status
        existing.tags = clean_tags
        existing.file_name = file.filename
        existing.file_url = file_url
        existing.document_text = resume_text
        existing.updated_at = datetime.utcnow()
        db.add(
            JobEvent(
                job_id=linked_job_id,
                owner_id=user_id,
                event_type=JobEventType.DOCUMENT,
                notes=f"Resume updated|Saved {existing.title}",
                created_at=datetime.utcnow(),
            )
        )
        db.commit()
        db.refresh(existing)
        return {
            "document_id": existing.id,
            "file_url": file_url,
            "title": existing.title,
            "status": existing.status,
            "tags": existing.tags,
            "job_id": existing.job_id,
            "version_label": existing.version_label,
            "version_number": existing.version_number,
        }

    record = Document(
        user_id=user_id,
        job_id=linked_job_id,
        title=clean_title or file.filename,
        doc_type=DocType.RESUME,
        status=document_status,
        tags=clean_tags,
        file_name=file.filename,
        file_url=file_url,
        document_text=resume_text,
        version_label="v1",
        version_number=1,
    )
    db.add(record)
    if linked_job_id:
        db.add(
            JobEvent(
                job_id=linked_job_id,
                owner_id=user_id,
                event_type=JobEventType.DOCUMENT,
                notes=f"Resume connected|Saved {record.title}",
                created_at=datetime.utcnow(),
            )
        )
    db.commit()
    db.refresh(record)

    return {
        "document_id": record.id,
        "file_url": file_url,
        "title": record.title,
        "status": record.status,
        "tags": record.tags,
        "job_id": record.job_id,
        "version_label": record.version_label,
        "version_number": record.version_number,
    }


@router.get("/resume/latest")
def get_latest_resume(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    user_id = current_user.get("sub")
    record = db.exec(
        select(Document)
        .where(Document.user_id == user_id)
        .where(Document.doc_type == DocType.RESUME)
        .order_by(Document.created_at.desc())
    ).first()

    if not record:
        raise HTTPException(status_code=404, detail="No resume found")

    return {
        "document_id": record.id,
        "file_url": record.file_url,
        "file_name": record.file_name,
        "resume_text": record.document_text,
    }


@router.post("/resume/parse-pdf")
async def parse_pdf(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
):
    if file.content_type != "application/pdf":
        raise HTTPException(status_code=400, detail="File must be a PDF")
    file_bytes = await file.read()
    try:
        import io

        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text_content = "".join(page.extract_text() or "" for page in reader.pages)
        if not text_content.strip():
            raise HTTPException(
                status_code=422,
                detail="Could not extract text from PDF. Try pasting text directly.",
            )  # noqa: E501
        return {"text": text_content.strip()}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF parsing failed: {str(e)}")


@router.post("/resume/generate-for-job")
async def generate_resume_for_job(
    payload: GenerateForJobRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    user_id = current_user.get("sub")
    job = db.exec(
        select(Job).where(Job.id == payload.job_id, Job.owner_id == user_id)
    ).first()  # noqa: E501
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    result = db.exec(
        text(
            "SELECT first_name, last_name, email, phone_number, professional_summary "
            "FROM users WHERE id = :user_id"
        ).bindparams(user_id=user_id)
    ).first()
    if not result:
        raise HTTPException(status_code=404, detail="User not found")

    profile_context = (
        f"Name: {result.first_name or ''} {result.last_name or ''}\n"
        f"Email: {result.email or ''}\n"
        f"Phone: {result.phone_number or 'Not provided'}\n"
        f"Summary: {result.professional_summary or 'Not provided'}"
    )
    job_context = (
        f"Company: {job.company}\nTitle: {job.title}\n"
        f"Job Posting:\n{job.job_posting_body}"
    )

    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1000,
        messages=[
            {
                "role": "user",
                "content": (
                    f"Generate a professional resume tailored for this job posting. "
                    f"Use the candidate profile to personalize it.\n\n"
                    f"CANDIDATE PROFILE:\n{profile_context}\n\n"
                    f"JOB DETAILS:\n{job_context}\n\n"
                    f"Return clean plain text resume only."
                ),
            }
        ],
    )
    return {"resume_text": message.content[0].text}


@router.post("/resume/improve")
async def improve_resume(
    payload: ImproveResumeRequest,
    current_user: dict = Depends(get_current_user),
):
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1000,
        messages=[
            {
                "role": "user",
                "content": (
                    f"Please improve the following resume based on this instruction: "
                    f"{payload.instruction}\n\n"
                    f"RESUME:\n{payload.resume_text}\n\n"
                    f"Return the improved resume in plain text only."
                ),
            }
        ],
    )

    return {"improved_text": message.content[0].text}


@router.get("/{document_id}/export")
def export_document(
    document_id: str,
    format: str,
    version_id: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Export a document as pdf, docx, or txt. Optionally export a
    specific past version instead of the current content.
    Rules: S3-BR-002 (ownership), S3-BR-001 (resume/cover_letter only), S3-BR-006"""
    if format not in EXPORT_CONTENT_TYPES:
        raise HTTPException(
            status_code=400, detail="Unsupported export format. Use pdf, docx, or txt."
        )

    user_id = current_user.get("sub")
    document = db.get(Document, document_id)
    if not document or document.user_id != user_id:
        raise HTTPException(status_code=404, detail="Document not found")

    text = document.document_text or ""
    if version_id:
        version = db.get(DocumentVersion, version_id)
        if (
            not version
            or version.document_id != document_id
            or version.user_id != user_id
        ):
            raise HTTPException(status_code=404, detail="Version not found")
        text = version.document_text or ""

    if not text.strip():
        raise HTTPException(
            status_code=400, detail="This document has no text content to export."
        )

    if format == "pdf":
        file_bytes = generate_pdf(text)
    elif format == "docx":
        file_bytes = generate_docx_bytes(text)
    else:
        file_bytes = text.encode("utf-8")

    filename = build_export_filename(document, format)

    return Response(
        content=file_bytes,
        media_type=EXPORT_CONTENT_TYPES[format],
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Cover letter endpoints ────────────────────────────────────────────────────


@router.post("/cover-letter/generate")
async def generate_cover_letter(
    payload: GenerateCoverLetterRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    user_id = current_user.get("sub")

    job = db.exec(
        select(Job).where(Job.id == payload.job_id, Job.owner_id == user_id)
    ).first()  # noqa: E501
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    skills = db.exec(
        select(Skill)
        .join(UserSkill, Skill.id == UserSkill.skill_id)
        .where(UserSkill.user_id == user_id)
    ).all()
    education = db.exec(
        select(Education).where(Education.user_id == user_id).order_by(Education.order)
    ).all()
    experiences = db.exec(
        select(Experience)
        .where(Experience.user_id == user_id)
        .order_by(Experience.order)
    ).all()

    skills_text = ", ".join([s.name for s in skills]) if skills else "Not provided"
    education_text = (
        "\n".join(
            [
                f"- {e.degree} in {e.field_of_study} at {e.school} ({e.start_date} – {e.end_date})"  # noqa: E501
                for e in education
            ]
        )
        if education
        else "Not provided"
    )
    experience_text = (
        "\n".join(
            [
                f"- {e.title} at {e.company} ({e.start_date} – {e.end_date}): {e.description}"  # noqa: E501
                for e in experiences
            ]
        )
        if experiences
        else "Not provided"
    )

    profile = payload.profile
    prompt = (
        f"Write a professional cover letter for the following job application.\n\n"
        f"Applicant:\n"
        f"- Name: {profile.get('firstName', '')} {profile.get('lastName', '')}\n"
        f"- Email: {profile.get('email', '')}\n"
        f"- Phone: {profile.get('phone', '')}\n"
        f"- Summary: {profile.get('summary', '')}\n\n"
        f"Skills: {skills_text}\n\n"
        f"Education:\n{education_text}\n\n"
        f"Work Experience:\n{experience_text}\n\n"
        f"Job:\n"
        f"- Title: {job.title}\n"
        f"- Company: {job.company}\n"
        f"- Location: {job.location or 'Not specified'}\n"
        f"- Description: {job.job_posting_body or 'Not provided'}\n\n"
        "Write a compelling, professional cover letter in first person. "
        "Tailor it specifically to the job description and highlight relevant "
        "skills and experience. Do not include a subject line or email headers. "
        "Start directly with the salutation."
    )

    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    message = client.messages.create(
        model="claude-opus-4-8",
        max_tokens=1024,
        messages=[{"role": "user", "content": prompt}],
    )

    return {"cover_letter": message.content[0].text}


@router.post("/cover-letter/improve")
async def improve_cover_letter(
    payload: ImproveCoverLetterRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    user_id = current_user.get("sub")

    job = db.exec(
        select(Job).where(Job.id == payload.job_id, Job.owner_id == user_id)
    ).first()  # noqa: E501
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    instruction = payload.instruction
    prompt = (
        "Improve the following cover letter based on this instruction: "
        f'"{instruction}"\n\n'
        f"Current cover letter:\n{payload.cover_letter_text}\n\n"
        "Return only the improved cover letter text, no explanations."
    )

    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    message = client.messages.create(
        model="claude-opus-4-8",
        max_tokens=1024,
        messages=[{"role": "user", "content": prompt}],
    )

    return {"improved_text": message.content[0].text}


@router.post("/cover-letter/save")
async def save_cover_letter(
    payload: SaveCoverLetterRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    user_id = current_user.get("sub")
    linked_job_id = payload.job_id or None

    job = None
    if linked_job_id:
        job = db.exec(
            select(Job).where(Job.id == linked_job_id, Job.owner_id == user_id)
        ).first()
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")

    document_status = parse_document_status(payload.status)
    clean_title = payload.title.strip() if payload.title else ""
    clean_tags = clean_document_tags(payload.tags)
    pdf_bytes = generate_pdf(payload.cover_letter_text)
    default_title = (
        f"Cover Letter - {job.company} {job.title}" if job else "Cover Letter"
    )
    document_title = clean_title or default_title
    file_name = f"{document_title}.pdf".replace(" ", "_")
    path = f"{user_id}/cover_letter/{int(time.time())}_{file_name}"

    file_url = await upload_to_storage(pdf_bytes, path, "application/pdf")

    existing = None
    if linked_job_id:
        existing = db.exec(
            select(Document)
            .where(Document.job_id == linked_job_id)
            .where(Document.user_id == user_id)
            .where(Document.doc_type == DocType.COVER_LETTER)
        ).first()

    if existing:
        snapshot_document_version(db, existing)
        existing.title = document_title
        existing.status = document_status
        existing.tags = clean_tags
        existing.document_text = payload.cover_letter_text
        existing.file_name = file_name
        existing.file_url = file_url
        existing.updated_at = datetime.utcnow()
        db.add(
            JobEvent(
                job_id=linked_job_id,
                owner_id=user_id,
                event_type=JobEventType.DOCUMENT,
                notes=f"Cover letter updated|Saved {document_title}",
                created_at=datetime.utcnow(),
            )
        )
        db.commit()
        db.refresh(existing)
        return {
            "document_id": existing.id,
            "file_url": file_url,
            "title": existing.title,
            "status": existing.status,
            "tags": existing.tags,
            "job_id": existing.job_id,
            "version_label": existing.version_label,
            "version_number": existing.version_number,
        }

    record = Document(
        user_id=user_id,
        job_id=linked_job_id,
        title=document_title,
        doc_type=DocType.COVER_LETTER,
        status=document_status,
        tags=clean_tags,
        file_name=file_name,
        file_url=file_url,
        document_text=payload.cover_letter_text,
        version_label="v1",
        version_number=1,
    )
    db.add(record)
    if linked_job_id:
        db.add(
            JobEvent(
                job_id=linked_job_id,
                owner_id=user_id,
                event_type=JobEventType.DOCUMENT,
                notes=f"Cover letter connected|Saved {document_title}",
                created_at=datetime.utcnow(),
            )
        )
    db.commit()
    db.refresh(record)
    return {
        "document_id": record.id,
        "file_url": file_url,
        "title": record.title,
        "status": record.status,
        "tags": record.tags,
        "job_id": record.job_id,
        "version_label": record.version_label,
        "version_number": record.version_number,
    }


@router.get("/cover-letter/job/{job_id}")
def get_cover_letter(
    job_id: str,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    user_id = current_user.get("sub")

    record = db.exec(
        select(Document)
        .where(Document.job_id == job_id)
        .where(Document.user_id == user_id)
        .where(Document.doc_type == DocType.COVER_LETTER)
    ).first()

    if not record:
        raise HTTPException(status_code=404, detail="No cover letter found")

    return {
        "document_id": record.id,
        "cover_letter_text": record.document_text,
        "file_url": record.file_url,
        "title": record.title,
        "status": record.status,
        "tags": record.tags,
        "version_label": record.version_label,
        "version_number": record.version_number,
    }


# ── General document endpoints ────────────────────────────────────────────────


@router.get("")
def get_documents(
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    user_id = current_user.get("sub")
    records = db.exec(
        select(Document)
        .where(Document.user_id == user_id)
        .order_by(Document.created_at.desc())
    ).all()
    return records


@router.post("/{document_id}/duplicate")
def duplicate_document(
    document_id: str,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    user_id = current_user.get("sub")
    source = db.get(Document, document_id)
    if not source or source.user_id != user_id:
        raise HTTPException(status_code=404, detail="Document not found")

    copy_title = (
        source.title if source.title.endswith("(Copy)") else f"{source.title} (Copy)"
    )
    duplicate = Document(
        user_id=user_id,
        job_id=None,
        title=copy_title,
        doc_type=source.doc_type,
        status=DocStatus.DRAFT,
        tags=source.tags,
        file_name=source.file_name,
        file_url=source.file_url,
        document_text=source.document_text,
        version_label="v1",
        version_number=1,
    )
    db.add(duplicate)
    db.commit()
    db.refresh(duplicate)
    return duplicate


@router.patch("/{document_id}/rename")
def rename_document(
    document_id: str,
    payload: RenameDocumentRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Rename a document. Metadata-only change; does not create a new
    version since the document's content is unchanged.
    Rules: S3-BR-007, S3-BR-008"""
    user_id = current_user.get("sub")
    document = db.get(Document, document_id)
    if not document or document.user_id != user_id:
        raise HTTPException(status_code=404, detail="Document not found")

    clean_title = payload.title.strip()
    if not clean_title:
        raise HTTPException(status_code=400, detail="Title cannot be empty")

    document.title = clean_title
    document.updated_at = datetime.utcnow()
    db.add(document)
    db.commit()
    db.refresh(document)
    return document


@router.patch("/{document_id}/link-to-job")
def link_document_to_job(
    document_id: str,
    payload: LinkDocumentRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Link a library document to a job. Enforces at most one resume and
    one cover letter linked per job at a time. Does not create a new
    version, only changes the job association.
    Rules: S3-BR-002, S3-BR-010, S3-BR-011, S3-BR-012"""
    user_id = current_user.get("sub")

    document = db.get(Document, document_id)
    if not document or document.user_id != user_id:
        raise HTTPException(status_code=404, detail="Document not found")

    if document.doc_type != payload.document_type:
        raise HTTPException(
            status_code=400,
            detail=(
                f"Document is a {DocType(document.doc_type).value}, "
                f"not a {DocType(payload.document_type).value}"
            ),
        )

    job = db.exec(
        select(Job).where(Job.id == payload.job_id, Job.owner_id == user_id)
    ).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    # S3-BR-010: at most one resume and one cover letter per job at a time.
    existing = db.exec(
        select(Document).where(
            Document.job_id == payload.job_id,
            Document.doc_type == payload.document_type,
            Document.user_id == user_id,
            Document.id != document_id,
        )
    ).first()

    if existing and not payload.replace_existing:
        # S3-BR-011: replacing a currently linked document requires confirmation.
        raise HTTPException(
            status_code=409,
            detail={
                "message": (
                    f"This job already has a linked {payload.document_type.value}. "
                    f"Set replace_existing=true to replace it."
                ),
                "requires_confirmation": True,
                "existing_document_id": existing.id,
            },
        )

    if existing and payload.replace_existing:
        existing.job_id = None
        existing.updated_at = datetime.utcnow()
        db.add(existing)

    document.job_id = payload.job_id
    document.updated_at = datetime.utcnow()
    db.add(document)
    db.commit()
    db.refresh(document)
    return document


@router.patch("/{document_id}/unlink-from-job")
def unlink_document_from_job(
    document_id: str,
    payload: UnlinkDocumentRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Unlink a document from a job without deleting it. Only succeeds
    if the document is currently linked to the given job.
    Rules: S3-BR-002, S3-BR-012"""
    user_id = current_user.get("sub")

    document = db.get(Document, document_id)
    if not document or document.user_id != user_id:
        raise HTTPException(status_code=404, detail="Document not found")

    if document.job_id != payload.job_id or document.doc_type != payload.document_type:
        raise HTTPException(
            status_code=400,
            detail="Document is not currently linked to this job.",
        )

    document.job_id = None
    document.updated_at = datetime.utcnow()
    db.add(document)
    db.commit()
    db.refresh(document)
    return document


@router.post("/{document_id}/archive")
def archive_document(
    document_id: str,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Archive a document without deleting it or its version history.
    Rules: S3-BR-009"""
    user_id = current_user.get("sub")
    document = db.get(Document, document_id)
    if not document or document.user_id != user_id:
        raise HTTPException(status_code=404, detail="Document not found")

    if document.status == DocStatus.ARCHIVED:
        raise HTTPException(status_code=400, detail="Document is already archived")

    document.status = DocStatus.ARCHIVED
    document.updated_at = datetime.utcnow()
    db.add(document)
    db.commit()
    db.refresh(document)
    return document


@router.post("/{document_id}/restore")
def restore_document(
    document_id: str,
    payload: RestoreDocumentRequest,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Restore an archived document to draft or final status.
    Rules: S3-BR-009"""
    user_id = current_user.get("sub")
    document = db.get(Document, document_id)
    if not document or document.user_id != user_id:
        raise HTTPException(status_code=404, detail="Document not found")

    if document.status != DocStatus.ARCHIVED:
        raise HTTPException(
            status_code=400, detail="Only archived documents can be restored"
        )

    restore_to = (payload.restore_to or "draft").lower()
    if restore_to == "archived":
        raise HTTPException(
            status_code=400, detail="Cannot restore a document to archived status"
        )

    status_map = {
        "draft": DocStatus.DRAFT,
        "active": DocStatus.DRAFT,
        "final": DocStatus.FINAL,
    }
    new_status = status_map.get(restore_to)
    if not new_status:
        raise HTTPException(status_code=400, detail="Invalid restore target status")

    document.status = new_status
    document.updated_at = datetime.utcnow()
    db.add(document)
    db.commit()
    db.refresh(document)
    return document


@router.delete("/{document_id}")
def delete_document(
    document_id: str,
    current_user: dict = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    user_id = current_user.get("sub")
    record = db.get(Document, document_id)
    if not record or record.user_id != user_id:
        raise HTTPException(status_code=404, detail="Document not found")
    db.delete(record)
    db.commit()
    return {"message": "Deleted"}
